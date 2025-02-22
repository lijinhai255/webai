import streamlit as st
import openai
import os
import pdfplumber
from docx import Document  # Use this instead of import docx
from dotenv import load_dotenv

# ✅ 1. 加载环境变量
load_dotenv()
openai_api_key = os.getenv("OPENAI_API_KEY")

# ✅ 2. 适配 openai>=1.0.0 版本（直接使用 openai.Client()）
client = openai.OpenAI(api_key=openai_api_key)

# ✅ 3. Streamlit UI
st.title("📄 AI 简历优化与生成工具")
import streamlit as st
import openai
import os
import pdfplumber
from docx import Document  # Use this instead of import docx
from dotenv import load_dotenv

# ✅ 1. 加载环境变量
load_dotenv()
openai_api_key = os.getenv("OPENAI_API_KEY")

# ✅ 2. 适配 openai>=1.0.0 版本（直接使用 openai.Client()）
client = openai.OpenAI(api_key=openai_api_key)

# ✅ 3. Streamlit UI
st.title("📄 AI 简历优化与生成工具")

# ✅ 4. 上传简历文件
uploaded_file = st.file_uploader("📂 上传你的简历文件（PDF / DOCX / TXT）", type=["pdf", "docx", "txt"])

# ✅ 5. 解析文件内容
def extract_text_from_file(file):
    text = ""
    if file.type == "application/pdf":
        with pdfplumber.open(file) as pdf:
            text = "\n".join(page.extract_text() for page in pdf.pages if page.extract_text())
    elif file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        doc = docx.Document(file)
        text = "\n".join([p.text for p in doc.paragraphs])
    elif file.type == "text/plain":
        text = file.read().decode("utf-8")
    return text

resume_text = ""
if uploaded_file:
    resume_text = extract_text_from_file(uploaded_file)
    st.text_area("📌 解析后的简历内容", resume_text, height=200)

# ✅ 6. 存储会话状态
if "ai_output" not in st.session_state:
    st.session_state.ai_output = ""
if "pending_message" not in st.session_state:
    st.session_state.pending_message = None  # 记录 AI 未完成的内容

# ✅ 7. 使用 OpenAI 进行简历优化
def analyze_resume_with_ai(resume_text, continue_from=None):
    prompt = f"""
你是一个专业的简历优化 AI 助手，目标是帮助用户提升简历质量，使其更具竞争力，并为面试做准备。

### **简历优化目标**
1. **优化简历内容**
2. **提供修改建议**
3. **面试准备**（列出知识点 & 预测面试问题）

### **用户简历**
{resume_text}

"""

    # ✅ 如果是“继续生成”，则加入之前的内容
    if continue_from:
        prompt += f"\n（继续上次的回答）\n{continue_from}"

    # ✅ 适配新版 openai>=1.0.0
    response = client.chat.completions.create(
        model="gpt-4",
        messages=[{"role": "system", "content": "你是一个专业的简历优化助手。"},
                  {"role": "user", "content": prompt}],
        max_tokens=1024,
        temperature=0.7
    )

    return response.choices[0].message.content  # ✅ 适配新版 API

# ✅ 8. 生成优化内容
if st.button("🚀 优化简历"):
    if not resume_text:
        st.warning("请输入简历信息或上传文件！")
    else:
        with st.spinner("AI 正在优化简历，请稍等..."):
            optimized_resume = analyze_resume_with_ai(resume_text)
            st.session_state.ai_output = optimized_resume  # ✅ 记录 AI 生成的内容
            st.session_state.pending_message = optimized_resume  # ✅ 记录未完成的部分

# ✅ 9. 显示 AI 生成的内容
if st.session_state.ai_output:
    sections = st.session_state.ai_output.split("---")
    if len(sections) >= 4:
        st.subheader("✅ **优化后的简历**")
        st.text_area("", sections[0].strip(), height=300)

        st.subheader("🔍 **修改建议**")
        st.markdown(sections[1].strip())

        st.subheader("🎯 **需要掌握的知识点**")
        st.markdown(sections[2].strip())

        st.subheader("💬 **可能遇到的面试问题**")
        st.markdown(sections[3].strip())

        # ✅ 继续生成按钮
        if st.button("继续生成"):
            with st.spinner("AI 继续生成中，请稍等..."):
                extra_output = analyze_resume_with_ai(resume_text, continue_from=st.session_state.pending_message)
                st.session_state.ai_output += "\n\n" + extra_output  # ✅ 追加到已有内容
                st.session_state.pending_message = extra_output  # ✅ 更新未完成的部分

        # ✅ 提供下载功能
        st.download_button("📥 下载优化简历", st.session_state.ai_output, file_name="optimized_resume.txt")

# ✅ 10. 版权信息
st.sidebar.markdown("---")
st.sidebar.markdown("💡 Created by hai | AI Resume Enhancer")

# ✅ 4. 上传简历文件
uploaded_file = st.file_uploader("📂 上传你的简历文件（PDF / DOCX / TXT）", type=["pdf", "docx", "txt"])

# ✅ 5. 解析文件内容
def extract_text_from_file(file):
    text = ""
    if file.type == "application/pdf":
        with pdfplumber.open(file) as pdf:
            text = "\n".join(page.extract_text() for page in pdf.pages if page.extract_text())
    elif file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        doc = docx.Document(file)
        text = "\n".join([p.text for p in doc.paragraphs])
    elif file.type == "text/plain":
        text = file.read().decode("utf-8")
    return text

resume_text = ""
if uploaded_file:
    resume_text = extract_text_from_file(uploaded_file)
    st.text_area("📌 解析后的简历内容", resume_text, height=200)

# ✅ 6. 存储会话状态
if "ai_output" not in st.session_state:
    st.session_state.ai_output = ""
if "pending_message" not in st.session_state:
    st.session_state.pending_message = None  # 记录 AI 未完成的内容

# ✅ 7. 使用 OpenAI 进行简历优化
def analyze_resume_with_ai(resume_text, continue_from=None):
    prompt = f"""
你是一个专业的简历优化 AI 助手，目标是帮助用户提升简历质量，使其更具竞争力，并为面试做准备。

### **简历优化目标**
1. **优化简历内容**
2. **提供修改建议**
3. **面试准备**（列出知识点 & 预测面试问题）

### **用户简历**
{resume_text}

"""

    # ✅ 如果是“继续生成”，则加入之前的内容
    if continue_from:
        prompt += f"\n（继续上次的回答）\n{continue_from}"

    # ✅ 适配新版 openai>=1.0.0
    response = client.chat.completions.create(
        model="gpt-4",
        messages=[{"role": "system", "content": "你是一个专业的简历优化助手。"},
                  {"role": "user", "content": prompt}],
        max_tokens=1024,
        temperature=0.7
    )

    return response.choices[0].message.content  # ✅ 适配新版 API

# ✅ 8. 生成优化内容
if st.button("🚀 优化简历"):
    if not resume_text:
        st.warning("请输入简历信息或上传文件！")
    else:
        with st.spinner("AI 正在优化简历，请稍等..."):
            optimized_resume = analyze_resume_with_ai(resume_text)
            st.session_state.ai_output = optimized_resume  # ✅ 记录 AI 生成的内容
            st.session_state.pending_message = optimized_resume  # ✅ 记录未完成的部分

# ✅ 9. 显示 AI 生成的内容
if st.session_state.ai_output:
    sections = st.session_state.ai_output.split("---")
    if len(sections) >= 4:
        st.subheader("✅ **优化后的简历**")
        st.text_area("", sections[0].strip(), height=300)

        st.subheader("🔍 **修改建议**")
        st.markdown(sections[1].strip())

        st.subheader("🎯 **需要掌握的知识点**")
        st.markdown(sections[2].strip())

        st.subheader("💬 **可能遇到的面试问题**")
        st.markdown(sections[3].strip())

        # ✅ 继续生成按钮
        if st.button("继续生成"):
            with st.spinner("AI 继续生成中，请稍等..."):
                extra_output = analyze_resume_with_ai(resume_text, continue_from=st.session_state.pending_message)
                st.session_state.ai_output += "\n\n" + extra_output  # ✅ 追加到已有内容
                st.session_state.pending_message = extra_output  # ✅ 更新未完成的部分

        # ✅ 提供下载功能
        st.download_button("📥 下载优化简历", st.session_state.ai_output, file_name="optimized_resume.txt")

# ✅ 10. 版权信息
st.sidebar.markdown("---")
st.sidebar.markdown("💡 Created by hai | AI Resume Enhancer")
